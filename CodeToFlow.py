import ast
import streamlit as st
import graphviz
import zipfile
import os
from io import BytesIO
from docx import Document
from docx.shared import Inches
from PIL import Image

# Function to parse Python code and extract control flow
def parse_code(code):
    tree = ast.parse(code)
    nodes = []
    edges = []

    # Traverse the AST (Abstract Syntax Tree)
    for node in ast.walk(tree):
        if isinstance(node, ast.FunctionDef):
            nodes.append((f"Function: {node.name}", "function"))
        elif isinstance(node, ast.If):
            nodes.append(("If Condition", "condition"))
        elif isinstance(node, ast.For):
            nodes.append(("For Loop", "loop"))
        elif isinstance(node, ast.While):
            nodes.append(("While Loop", "loop"))
        elif isinstance(node, ast.Call):
            if isinstance(node.func, ast.Attribute):
                nodes.append((f"API Call: {node.func.attr}", "api"))
            elif isinstance(node.func, ast.Name):
                nodes.append((f"Function Call: {node.func.id}", "call"))

    return nodes

# Function to generate a flow diagram using Graphviz
def generate_flow_diagram(nodes):
    dot = graphviz.Digraph()
    for i, (node_name, node_type) in enumerate(nodes):
        dot.node(str(i), node_name, shape="box" if node_type == "function" else "ellipse")
        if i > 0:
            dot.edge(str(i - 1), str(i))
    return dot

# Function to create a Word document with the flow diagram and description
def create_word_document(diagram_image_path, description, code_snippets):
    doc = Document()
    doc.add_heading("Code Flow Diagram and Documentation", level=1)

    # Add the flow diagram image
    doc.add_heading("Flow Diagram", level=2)
    doc.add_picture(diagram_image_path, width=Inches(6))

    # Add the description
    doc.add_heading("Flow Description", level=2)
    doc.add_paragraph(description)

    # Add code snippets
    doc.add_heading("Code Snippets", level=2)
    for snippet in code_snippets:
        doc.add_paragraph(snippet)

    return doc

# Function to process uploaded files
def process_files(uploaded_files):
    all_nodes = []
    code_snippets = []
    for uploaded_file in uploaded_files:
        code = uploaded_file.read().decode("utf-8")
        nodes = parse_code(code)
        all_nodes.extend(nodes)
        code_snippets.append(code)
    return all_nodes, code_snippets

# Streamlit App
st.title("Code Flow Diagram Generator with Documentation")

# File uploader
uploaded_files = st.file_uploader("Upload Python files or a zip folder", type=["py", "zip"], accept_multiple_files=True)

if uploaded_files:
    # Check if a zip file is uploaded
    if any(file.name.endswith(".zip") for file in uploaded_files):
        with zipfile.ZipFile(uploaded_files[0], "r") as zip_ref:
            zip_ref.extractall("temp")
            uploaded_files = [open(os.path.join("temp", f), "rb") for f in os.listdir("temp")]

    # Process uploaded files
    nodes, code_snippets = process_files(uploaded_files)

    # Generate and display the flow diagram
    if nodes:
        st.subheader("Generated Flow Diagram")
        dot = generate_flow_diagram(nodes)

        # Save the diagram as an image
        diagram_image_path = "flow_diagram.png"

        dot.render("flow_diagram", format="png", cleanup=True)

        # Display the diagram in Streamlit
        try:
            st.image(diagram_image_path)
        except FileNotFoundError:
            st.error(f"File not found: {diagram_image_path}. Ensure the file is saved correctly.")

        # Create a description of the flow
        description = "This is a flow diagram generated from the provided Python code. It includes functions, loops, conditions, and API calls."

        # Create a Word document
        doc = create_word_document(diagram_image_path, description, code_snippets)

        # Save the Word document to a BytesIO object
        doc_bytes = BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)

        # Add a download button for the Word document
        st.download_button(
            label="Download Word Document",
            data=doc_bytes,
            file_name="code_flow_documentation.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    else:
        st.error("No valid Python code found in the uploaded files.")