import ast
import streamlit as st
import graphviz
import zipfile
import os
import shutil
from io import BytesIO
from docx import Document
from docx.shared import Inches
from PIL import Image
import time

# Create an images folder if it doesn't exist
if not os.path.exists("images"):
    os.makedirs("images")

# Function to parse Python code and extract functions/classes
def parse_code(code):
    tree = ast.parse(code)
    functions = []
    classes = []

    # Traverse the AST (Abstract Syntax Tree)
    for node in ast.walk(tree):
        if isinstance(node, ast.FunctionDef):
            functions.append(node)
        elif isinstance(node, ast.ClassDef):
            classes.append(node)

    return functions, classes

# Function to extract nodes (control flow elements) from a function
def parse_function(func):
    nodes = []
    for node in ast.walk(func):
        if isinstance(node, ast.FunctionDef):
            nodes.append((f"Function: {node.name}", "function"))
        elif isinstance(node, ast.If):
            # Extract the actual condition
            condition = ast.unparse(node.test).strip()  # Get the condition as a string
            nodes.append((f"If Condition: {condition}", "condition"))
        elif isinstance(node, ast.For):
            nodes.append(("For Loop", "loop"))
        elif isinstance(node, ast.While):
            # Extract the actual condition
            condition = ast.unparse(node.test).strip()  # Get the condition as a string
            nodes.append((f"While Loop: {condition}", "loop"))
        elif isinstance(node, ast.Call):
            if isinstance(node.func, ast.Attribute):
                nodes.append((f"API Call: {node.func.attr}", "api"))
            elif isinstance(node.func, ast.Name):
                nodes.append((f"Function Call: {node.func.id}", "call"))
    return nodes

# Function to extract nodes (control flow elements) from a class
def parse_class(cls):
    nodes = []
    for node in ast.walk(cls):
        if isinstance(node, ast.FunctionDef):
            nodes.append((f"Method: {node.name}", "function"))
        elif isinstance(node, ast.If):
            # Extract the actual condition
            condition = ast.unparse(node.test).strip()  # Get the condition as a string
            nodes.append((f"If Condition: {condition}", "condition"))
        elif isinstance(node, ast.For):
            nodes.append(("For Loop", "loop"))
        elif isinstance(node, ast.While):
            # Extract the actual condition
            condition = ast.unparse(node.test).strip()  # Get the condition as a string
            nodes.append((f"While Loop: {condition}", "loop"))
        elif isinstance(node, ast.Call):
            if isinstance(node.func, ast.Attribute):
                nodes.append((f"API Call: {node.func.attr}", "api"))
            elif isinstance(node.func, ast.Name):
                nodes.append((f"Function Call: {node.func.id}", "call"))
    return nodes

# Function to generate a flow diagram for a function or class
def generate_flow_diagram(name, nodes):
    dot = graphviz.Digraph(engine="dot")
    dot.attr(rankdir="TB", size="10", dpi="300")  # Adjust size and DPI for better clarity

    # Add nodes with different styles based on type
    for i, (node_name, node_type) in enumerate(nodes):
        if node_type == "function":
            dot.node(str(i), node_name, shape="box", style="filled", fillcolor="lightblue")
        elif node_type == "condition":
            dot.node(str(i), node_name, shape="diamond", style="filled", fillcolor="lightyellow")
        elif node_type == "loop":
            dot.node(str(i), node_name, shape="ellipse", style="filled", fillcolor="lightgreen")
        elif node_type == "api":
            dot.node(str(i), node_name, shape="oval", style="filled", fillcolor="orange")
        elif node_type == "call":
            dot.node(str(i), node_name, shape="ellipse", style="filled", fillcolor="lightgray")

    # Add edges between nodes
    for i in range(1, len(nodes)):
        dot.edge(str(i - 1), str(i))

    return dot

# Function to create a Word document with function/class-wise diagrams and explanations
def create_word_document(functions, classes, code_snippets):
    doc = Document()
    doc.add_heading("Code Flow Diagram and Documentation", level=1)

    # Add diagrams and explanations for functions
    if functions:
        doc.add_heading("Functions", level=2)
        for func in functions:
            # Generate diagram for the function
            nodes = parse_function(func)
            dot = generate_flow_diagram(func.name, nodes)

            # Save the diagram as an image in the images folder
            diagram_image_path = os.path.join("images", f"{func.name}_flow.png")
            dot.render(os.path.join("images", f"{func.name}_flow"), format="png", cleanup=True)

            # Resize the image for better fit in the Word document
            with Image.open(diagram_image_path) as img:
                img = img.resize((800, 600))  # Resize to fit the document
                img.save(diagram_image_path)

            # Add the diagram to the Word document
            doc.add_heading(f"Function: {func.name}", level=3)
            doc.add_picture(diagram_image_path, width=Inches(6))

            # Add explanation and code snippet
            doc.add_paragraph(f"Explanation for function '{func.name}':")
            doc.add_paragraph(ast.get_docstring(func) or "No docstring available. This function performs a specific task.")
            doc.add_paragraph("Code Snippet:")
            try:
                # Use ast.unparse() if available (Python 3.9+)
                doc.add_paragraph(ast.unparse(func))
            except AttributeError:
                # Fallback to ast.dump() for older Python versions
                doc.add_paragraph(ast.dump(func))

    # Add diagrams and explanations for classes
    if classes:
        doc.add_heading("Classes", level=2)
        for cls in classes:
            # Generate diagram for the class
            nodes = parse_class(cls)
            dot = generate_flow_diagram(cls.name, nodes)

            # Save the diagram as an image in the images folder
            diagram_image_path = os.path.join("images", f"{cls.name}_flow.png")
            dot.render(os.path.join("images", f"{cls.name}_flow"), format="png", cleanup=True)

            # Resize the image for better fit in the Word document
            with Image.open(diagram_image_path) as img:
                img = img.resize((800, 800))  # Resize to fit the document
                img.save(diagram_image_path)

            # Add the diagram to the Word document
            doc.add_heading(f"Class: {cls.name}", level=3)
            doc.add_picture(diagram_image_path, width=Inches(6))

            # Add explanation and code snippet
            doc.add_paragraph(f"Explanation for class '{cls.name}':")
            doc.add_paragraph(ast.get_docstring(cls) or "No docstring available. This class represents a specific entity or functionality.")
            doc.add_paragraph("Code Snippet:")
            try:
                # Use ast.unparse() if available (Python 3.9+)
                doc.add_paragraph(ast.unparse(cls))
            except AttributeError:
                # Fallback to ast.dump() for older Python versions
                doc.add_paragraph(ast.dump(cls))

    return doc

# Function to process uploaded files
def process_files(uploaded_files):
    functions = []
    classes = []
    code_snippets = []
    for uploaded_file in uploaded_files:
        code = uploaded_file.read().decode("utf-8")
        funcs, cls = parse_code(code)
        functions.extend(funcs)
        classes.extend(cls)
        code_snippets.append(code)
    return functions, classes, code_snippets

# Streamlit App
st.title("Code Flow Diagram Generator with Documentation")

# File uploader
uploaded_files = st.file_uploader("Upload Python files or a zip folder", type=["py", "zip"], accept_multiple_files=True)

# Function to recursively find all .py files in a directory
def find_py_files(directory):
    py_files = []
    for root, _, files in os.walk(directory):
        for file in files:
            if file.endswith(".py"):
                py_files.append(os.path.join(root, file))
    return py_files

# Function to safely delete a directory
def delete_directory(directory):
    try:
        shutil.rmtree(directory)
    except OSError as e:
        if e.errno == 39:  # Directory not empty
            print(f"Directory not empty: {directory}. Retrying after a short delay...")
            time.sleep(1)  # Wait for 1 second
            shutil.rmtree(directory, ignore_errors=True)  # Force delete
        else:
            raise e

# Updated file processing logic
if uploaded_files:
    # Check if a zip file is uploaded
    if any(file.name.endswith(".zip") for file in uploaded_files):
        with zipfile.ZipFile(uploaded_files[0], "r") as zip_ref:
            zip_ref.extractall("temp")

        # Find all .py files in the extracted directory (including nested folders)
        py_files = find_py_files("temp")

        if not py_files:
            st.error("No Python files (.py) found in the uploaded zip file.")
        else:
            # Open all .py files for processing
            uploaded_files = [open(py_file, "rb") for py_file in py_files]

            # Process uploaded files
            functions, classes, code_snippets = process_files(uploaded_files)

            # Close all opened files
            for file in uploaded_files:
                file.close()

            # Generate and display the flow diagrams
            if functions or classes:
                st.subheader("Generated Flow Diagrams")

                # Create a Word document
                doc = create_word_document(functions, classes, code_snippets)

                # Save the Word document to a BytesIO object
                doc_bytes = BytesIO()
                doc.save(doc_bytes)
                doc_bytes.seek(0)

                # Add a download button for the Word document
                st.download_button(
                    label="Download Word Document",
                    data=doc_bytes,
                    file_name="code_flow_documentation.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            else:
                st.error("No valid Python code found in the uploaded files.")
    else:
        # Process non-zip files (individual .py files)
        uploaded_files = [file for file in uploaded_files if file.name.endswith(".py")]
        if not uploaded_files:
            st.error("No Python files (.py) found in the uploaded files.")
        else:
            # Process uploaded files
            functions, classes, code_snippets = process_files(uploaded_files)

            # Generate and display the flow diagrams
            if functions or classes:
                st.subheader("Generated Flow Diagrams")

                # Create a Word document
                doc = create_word_document(functions, classes, code_snippets)

                # Save the Word document to a BytesIO object
                doc_bytes = BytesIO()
                doc.save(doc_bytes)
                doc_bytes.seek(0)

                # Add a download button for the Word document
                st.download_button(
                    label="Download Word Document",
                    data=doc_bytes,
                    file_name="code_flow_documentation.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            else:
                st.error("No valid Python code found in the uploaded files.")

    # Clean up temporary files
    if os.path.exists("temp"):
        delete_directory("temp")