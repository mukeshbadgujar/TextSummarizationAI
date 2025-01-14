from langchain.text_splitter import CharacterTextSplitter
from langchain.prompts import PromptTemplate
from langchain.llms import HuggingFacePipeline
from langchain.schema import Document
from transformers import pipeline
import streamlit as st
import torch
import PyPDF2
from docx import Document as DocxDocument
from io import StringIO

# Initialize the summarization model
@st.cache_resource
def load_summarizer():
    device = 0 if torch.cuda.is_available() else -1
    device = -1  # Force CPU
    summarizer_pipeline = pipeline("summarization", model="facebook/bart-large-cnn", device=device)
    return HuggingFacePipeline(pipeline=summarizer_pipeline)

# Initialize the question-answering model
@st.cache_resource
def load_qa_model():
    device = 0 if torch.cuda.is_available() else -1
    device = -1  # Force CPU
    qa_pipeline = pipeline("question-answering", model="distilbert-base-uncased-distilled-squad", device=device)
    return qa_pipeline

# Function to extract text from PDF
def extract_text_from_pdf(file):
    pdf_reader = PyPDF2.PdfReader(file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text()
    return text

# Function to extract text from Word
def extract_text_from_word(file):
    doc = DocxDocument(file)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

# Load models
summarizer = load_summarizer()
qa_model = load_qa_model()

# Streamlit App
st.title("Summarize and Extract Answers")

uploaded_file = st.file_uploader("Upload your file", type=["txt", "md", "pdf", "docx"])

if uploaded_file is not None:
    # Extract text based on file type
    if uploaded_file.type == "application/pdf":
        text_data = extract_text_from_pdf(uploaded_file)
    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        text_data = extract_text_from_word(uploaded_file)
    else:
        text_data = uploaded_file.read().decode("utf-8").strip()

    if not text_data:
        st.error("The uploaded file is empty or could not be processed. Please upload a valid file.")
    else:
        st.subheader("Uploaded Text")
        st.text_area("File Content", text_data, height=200)

        # Split the text into chunks
        text_splitter = CharacterTextSplitter(separator="\n", chunk_size=1000, chunk_overlap=100)
        chunks = text_splitter.split_text(text_data)
        documents = [Document(page_content=chunk) for chunk in chunks]

        st.write(f"Document split into {len(documents)} chunks.")

        # Summarize the data
        if st.button("Summarize"):
            if not chunks:
                st.error("The text file is too small to summarize.")
            else:
                st.info("Summarizing... Please wait.")
                try:
                    summaries = []  # To store individual summaries
                    for i, chunk in enumerate(chunks):
                        st.info(f"Summarizing chunk {i + 1}/{len(chunks)}...")
                        try:
                            summary_output = summarizer(chunk)
                            summaries.append(summary_output)
                        except Exception as chunk_error:
                            st.error(f"Error summarizing chunk {i + 1}: {chunk_error}")
                            break

                    # Combine summaries if all chunks are processed successfully
                    if summaries:
                        final_summary = " ".join(summaries)
                        st.subheader("Summary")
                        st.write(final_summary)
                    else:
                        st.error("No summaries could be generated.")

                except Exception as e:
                    st.error(f"Error during summarization: {e}")

        # Ask a question about the text
        user_question = st.text_input("Ask a question about the text")
        if user_question:
            if not chunks:
                st.error("The text file is too small to extract answers.")
            else:
                st.info("Searching for an answer...")
                try:
                    answers = []

                    # Process each chunk
                    for i, chunk in enumerate(chunks):
                        st.info(f"Processing chunk {i + 1}/{len(chunks)}...")
                        try:
                            result = qa_model({"question": user_question, "context": chunk})
                            answers.append(result)
                        except Exception as chunk_error:
                            st.warning(f"Error processing chunk {i + 1}: {chunk_error}")
                            continue

                    # Combine and refine answers
                    if answers:
                        # Rank answers by score
                        ranked_answers = sorted(answers, key=lambda x: x.get("score", 0), reverse=True)

                        # Display the best answer
                        best_answer = ranked_answers[0]
                        st.subheader("Answer")
                        st.write(best_answer["answer"])
                        st.write(f"Confidence: {best_answer['score']:.2f}")

                        # Show alternative answers (if needed)
                        st.subheader("Alternative Answers")
                        for alt in ranked_answers[1:3]:  # Show top 3 answers
                            st.write(f"- {alt['answer']} (Confidence: {alt['score']:.2f})")
                    else:
                        st.error("No answers could be generated from the text.")

                except Exception as e:
                    st.error(f"Error during question answering: {e}")